"""Build the final Word report from docs/final_report.md.

Design contract: ``standard_business_brief`` with one consistent Korean-font
override (Malgun Gothic for East Asian glyphs). The user-requested document
furniture override removes all running headers, footers, rules, and page
numbers. Tables use fixed 9360-DXA geometry.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
BASE_FONT = "Calibri"
EAST_ASIA_FONT = "맑은 고딕"
HEADING_BLUE = "2E74B5"
HEADING_DARK_BLUE = "1F4D78"
TABLE_HEADER_FILL = "F2F4F7"


def set_font(run, *, name: str = BASE_FONT, east_asia: str = EAST_ASIA_FONT, size: float | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:cs"), name)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_style_font(style, *, name: str, east_asia: str, size: float, color: str = "000000", bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:cs"), name)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, name=BASE_FONT, east_asia=EAST_ASIA_FONT, size=11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    specs = {
        "Heading 1": (16, HEADING_BLUE, 16, 8),
        "Heading 2": (13, HEADING_BLUE, 12, 6),
        "Heading 3": (12, HEADING_DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in specs.items():
        style = doc.styles[style_name]
        set_style_font(style, name=BASE_FONT, east_asia=EAST_ASIA_FONT, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def configure_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    # Explicitly keep all page furniture empty.
    for part in (section.header, section.footer):
        for paragraph in part.paragraphs:
            for run in list(paragraph.runs):
                paragraph._p.remove(run._r)


def add_inline_runs(paragraph, text: str, *, size: float | None = None, bold: bool = False) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_font(run, size=size)
            run.bold = bold
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Consolas", east_asia=EAST_ASIA_FONT, size=size or 9.5)
            run.font.color.rgb = RGBColor.from_string(HEADING_DARK_BLUE)
        else:
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=size)
            run.bold = True
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_font(run, size=size)
        run.bold = bold


def add_numbering_definition(doc: Document, kind: str, marker_text: str | None = None) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{0xA1000000 + abstract_id:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    template = OxmlElement("w:tmpl")
    template.set(qn("w:val"), f"{0xB1000000 + abstract_id:08X}")
    abstract.append(template)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    if kind == "bullet":
        num_format = "bullet"
        level_marker = "•"
    else:
        num_format = "decimal"
        level_marker = "%1."
    num_fmt.set(qn("w:val"), num_format)
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), level_marker)
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BASE_FONT)
    fonts.set(qn("w:hAnsi"), BASE_FONT)
    fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), marker_text or "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def add_list_paragraph(doc: Document, text: str, *, num_id: int) -> None:
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    paragraph.paragraph_format.keep_together = True
    add_inline_runs(paragraph, text)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGINS_DXA.items():
        tag = qn(f"w:{edge}")
        node = tc_mar.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "B7C2CE")
        borders.append(node)


def table_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    if cols == 1:
        return [9360]
    if cols == 2:
        first_max = max(len(row[0]) for row in rows)
        if first_max <= 14:
            return [2700, 6660]
        return [3600, 5760]
    if cols == 3:
        return [2340, 3510, 3510]
    if cols == 4:
        return [1800, 2520, 2520, 2520]
    base, remainder = divmod(CONTENT_WIDTH_DXA, cols)
    return [base + (1 if idx < remainder else 0) for idx in range(cols)]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    widths = table_widths(rows)
    table = doc.add_table(rows=len(rows), cols=len(widths))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    set_table_borders(table)

    table_font_size = 8.0 if len(widths) >= 6 else 9.0
    for row_idx, (row, values) in enumerate(zip(table.rows, rows)):
        row_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_pr.append(cant_split)
        if row_idx == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            row_pr.append(repeat)
        for col_idx, (cell, value, width) in enumerate(zip(row.cells, values, widths)):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), TABLE_HEADER_FILL)
                tc_pr.append(shd)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline_runs(paragraph, value, size=table_font_size, bold=row_idx == 0)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(2)


def add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shading)
    run = paragraph.add_run("\n".join(lines))
    set_font(run, name="Consolas", east_asia=EAST_ASIA_FONT, size=8.5)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw_rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        raw_rows.append(cells)
        index += 1
    rows = [
        row
        for row in raw_rows
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in row)
    ]
    return rows, index


def add_cover(doc: Document, lines: list[str]) -> None:
    title = lines[0].lstrip("# ").strip()
    subtitle = next(line.lstrip("# ").strip() for line in lines[1:] if line.startswith("## "))
    table_start = next(idx for idx, line in enumerate(lines) if line.strip().startswith("|"))
    metadata, table_end = parse_table(lines, table_start)
    description = " ".join(line.strip() for line in lines[table_end:] if line.strip())

    top = doc.add_paragraph()
    top.paragraph_format.space_after = Pt(54)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_after = Pt(14)
    title_p.paragraph_format.keep_with_next = True
    title_run = title_p.add_run(title)
    set_font(title_run, name=BASE_FONT, east_asia=EAST_ASIA_FONT, size=22)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(HEADING_DARK_BLUE)

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(30)
    subtitle_run = subtitle_p.add_run(subtitle)
    set_font(subtitle_run, size=15)
    subtitle_run.bold = True
    subtitle_run.font.color.rgb = RGBColor.from_string(HEADING_BLUE)

    add_table(doc, metadata)

    desc_p = doc.add_paragraph()
    desc_p.paragraph_format.space_before = Pt(18)
    desc_p.paragraph_format.space_after = Pt(0)
    desc_p.paragraph_format.line_spacing = 1.2
    add_inline_runs(desc_p, description, size=10.5)


def build_docx(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_page(section)
    doc.settings.odd_and_even_pages_header_footer = False

    props = doc.core_properties
    props.title = "의미 기반 미디어 전송 검증시스템 및 멀티모달 시맨틱 미디어 구조 연구"
    props.subject = "용역결과보고서"
    props.keywords = "ETRI, 시맨틱 미디어, QVHighlights, 검증시스템"
    props.comments = "v1.1.0 최종 납품본"

    first_break = lines.index("<!-- PAGEBREAK -->")
    add_cover(doc, lines[:first_break])
    doc.add_page_break()

    index = first_break + 1
    in_code = False
    code_lines: list[str] = []
    current_list_kind: str | None = None
    current_num_id: int | None = None

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if in_code:
            if stripped.startswith("```"):
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                code_lines.append(raw)
            index += 1
            continue

        if stripped.startswith("```"):
            in_code = True
            current_list_kind = None
            current_num_id = None
            index += 1
            continue
        if stripped == "<!-- PAGEBREAK -->":
            doc.add_page_break()
            current_list_kind = None
            current_num_id = None
            index += 1
            continue
        if not stripped:
            current_list_kind = None
            current_num_id = None
            index += 1
            continue
        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            current_list_kind = None
            current_num_id = None
            continue
        if stripped.startswith("### "):
            paragraph = doc.add_heading(stripped[4:].strip(), level=3)
            for run in paragraph.runs:
                set_font(run, size=12)
            current_list_kind = None
            current_num_id = None
            index += 1
            continue
        if stripped.startswith("## "):
            paragraph = doc.add_heading(stripped[3:].strip(), level=2)
            for run in paragraph.runs:
                set_font(run, size=13)
            current_list_kind = None
            current_num_id = None
            index += 1
            continue
        if stripped.startswith("# "):
            paragraph = doc.add_heading(stripped[2:].strip(), level=1)
            for run in paragraph.runs:
                set_font(run, size=16)
            current_list_kind = None
            current_num_id = None
            index += 1
            continue

        ordered = re.match(r"^\d+\.\s+(.*)$", stripped)
        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if ordered or bullet:
            kind = "decimal" if ordered else "bullet"
            text = (ordered or bullet).group(1)
            if current_list_kind != kind or current_num_id is None:
                current_list_kind = kind
                current_num_id = add_numbering_definition(doc, kind)
            add_list_paragraph(doc, text, num_id=current_num_id)
            index += 1
            continue

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.keep_together = False
        add_inline_runs(paragraph, stripped)
        current_list_kind = None
        current_num_id = None
        index += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build the final Word result report.")
    parser.add_argument("--source", type=Path, default=Path("docs/final_report.md"))
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("deliverables/ETRI_시맨틱_미디어_용역결과보고서.docx"),
    )
    args = parser.parse_args()
    build_docx(args.source, args.output)
    print(args.output.resolve())


if __name__ == "__main__":
    main()
