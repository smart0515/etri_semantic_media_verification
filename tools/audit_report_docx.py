"""Audit the final Word report's structure and delivery-specific layout rules."""

from __future__ import annotations

import argparse
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


CONTENT_WIDTH_DXA = 9360


def require(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def int_attr(element, name: str) -> int:
    require(element is not None, f"missing {name}")
    return int(element.get(qn(name)))


def audit(path: Path) -> None:
    doc = Document(path)
    require(len(doc.sections) == 1, "report must use one section")
    section = doc.sections[0]
    require(section.page_width.inches == 8.5, "page width must be Letter 8.5 in")
    require(section.page_height.inches == 11.0, "page height must be Letter 11 in")
    for name in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        require(abs(getattr(section, name).inches - 1.0) < 0.001, f"{name} must be 1 in")

    header_text = "".join(node.text or "" for node in section.header._element.iter(qn("w:t"))).strip()
    footer_text = "".join(node.text or "" for node in section.footer._element.iter(qn("w:t"))).strip()
    require(not header_text, "running header must be empty")
    require(not footer_text, "running footer must be empty")

    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    require("최종 요약" in paragraphs, "final summary heading is missing")
    require(paragraphs.index("최종 요약") > paragraphs.index("부록 C. 참고자료"), "summary must follow appendices")
    require(paragraphs[-1].startswith("최종 납품물은 QVHighlights 호환 데이터를 읽고"), "summary must be the report's final content")

    for table_index, table in enumerate(doc.tables, start=1):
        tbl_pr = table._tbl.tblPr
        require(int_attr(tbl_pr.find(qn("w:tblW")), "w:w") == CONTENT_WIDTH_DXA, f"table {table_index}: wrong width")
        require(int_attr(tbl_pr.find(qn("w:tblInd")), "w:w") == 120, f"table {table_index}: wrong indent")
        grid_widths = [int_attr(col, "w:w") for col in table._tbl.tblGrid.findall(qn("w:gridCol"))]
        require(sum(grid_widths) == CONTENT_WIDTH_DXA, f"table {table_index}: grid width mismatch")
        for row_index, row in enumerate(table.rows, start=1):
            cell_widths = [int_attr(cell._tc.get_or_add_tcPr().find(qn("w:tcW")), "w:w") for cell in row.cells]
            require(cell_widths == grid_widths, f"table {table_index} row {row_index}: cell/grid mismatch")

    with ZipFile(path) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
        require("PAGE" not in document_xml, "page-number field is not allowed")
        require("상 시스템의 실제 성능을 의미하지 않는다" not in document_xml, "draft disclaimer remains")
        require(document_xml.count("<w:numPr>") > 0, "real Word numbering definitions are required")

    print(f"PASS: {path.resolve()}")
    print(f"tables={len(doc.tables)}, body_paragraphs={len(doc.paragraphs)}, summary_last=yes, page_furniture=empty")


def main() -> None:
    parser = argparse.ArgumentParser(description="Audit the final Word report.")
    parser.add_argument(
        "path",
        nargs="?",
        type=Path,
        default=Path("deliverables/ETRI_시맨틱_미디어_용역결과보고서.docx"),
    )
    args = parser.parse_args()
    audit(args.path)


if __name__ == "__main__":
    main()
